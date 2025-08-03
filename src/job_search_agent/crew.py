from crewai import Agent, Crew, Process, Task, LLM
from crewai.project import CrewBase, agent, crew, task, before_kickoff, after_kickoff
from crewai.agents.agent_builder.base_agent import BaseAgent
from crewai_tools import FileReadTool, FileWriterTool, DirectoryReadTool, SerperDevTool, ScrapeWebsiteTool
from job_search_agent.schemas import Resume, JobPostings, TailoredCVs, TailoredCoverLetters

from typing import List
from dotenv import load_dotenv
import docx, pdfplumber
import os, json, re

# -- Load environment variables --
load_dotenv()
GEMINI_MODEL = os.environ.get("MODEL")
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
SERPER_API_KEY = os.environ.get("SERPER_API_KEY")
MAX_RPM = 20

# -- Create LLM client --
gemini_llm = LLM(
    model=GEMINI_MODEL,
    api_key=GEMINI_API_KEY,
    temperature=0.5
)

# -- Initialize tools --
file_read_tool = FileReadTool()
file_write_tool = FileWriterTool()
directory_read_tool = DirectoryReadTool()
search_tool = SerperDevTool(
    country="vn",
    locale="vn",
    location="Hanoi, Hanoi, Vietnam",
    n_results=20
)
scrape_tool = ScrapeWebsiteTool()

# -- Define a utility function to sanitize text --
def sanitize(text):
    text = re.sub(r'[\\/*?:"<>|]', "", text)  # remove invalid chars
    text = text.strip().replace(" ", "_")     # optional: replace spaces
    return text

# -- Define the JobSearchAgent crew --
@CrewBase
class JobSearchAgent():
    """JobSearchAgent crew"""

    agents: List[BaseAgent]
    tasks: List[Task]

    # -- Before kickoff function --
    @before_kickoff
    def prepare_input(self, inputs):
        """
        Verify the input file path exists and is accessible.
        If valid, extract the content and return it in plain text format.
        If invalid, raise an error.
        """
        input_folder = inputs.get("input_folder", "input")
        print(f"Checking input from folder: {input_folder}")
        allowed_ext = [".pdf", ".docx", ".txt", ".md"]
        resume_content = ""
        resume_found = False

        for filename in os.listdir(input_folder):
            ext = os.path.splitext(filename)[1].lower()
            file_path = os.path.join(input_folder, filename)
            if os.path.isfile(file_path) and ext in allowed_ext:
                resume_found = True
                print(f"\n📃 Parsing file: {file_path}")

                try:
                    if ext == ".pdf":
                        with pdfplumber.open(file_path) as pdf:
                            for page in pdf.pages:
                                resume_content += page.extract_text() + "\n"
                    elif ext == ".docx":
                        doc = docx.Document(file_path)
                        for para in doc.paragraphs:
                            resume_content += para.text + "\n"
                    elif ext in [".txt", ".md"]:
                        with open(file_path, "r", encoding="utf-8") as file:
                            resume_content += file.read() + "\n"
                    break # Stop after the first valid file
                except Exception as e:
                    raise RuntimeError(f"Error parsing file {file_path}: {str(e)}")
        if not resume_found:
            raise FileNotFoundError(f"No valid resume files found in {input_folder}. Supported formats: {', '.join(allowed_ext)}")
        if not resume_content.strip():
            raise ValueError("The resume content is empty. Please provide a valid resume file with content.")
        
        print("Resume content extracted successfully.")

        # Add resume content to the input
        inputs["resume_content"] = resume_content.strip()
        return inputs

    # -- After kickoff function --
    @after_kickoff
    def organize_output_files(self, output):
        """
        Read output/tailored_cv.json & output/tailored_cover_letter.json files.
        Separate the content by company and job title into subdirectories named '{company}-{title}'.
        Each subdirectory will contain a 'cv.docx' and 'cover_letter.docx' file matching the company and job title.
        """
        print("\n🗂️ Organizing output files...")
        
        output_folder = "output"

        # Ensure output folder exists
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        cvs_file_path = os.path.join(output_folder, "tailored_cv.json")
        cover_letters_file_path = os.path.join(output_folder, "tailored_cover_letter.json")

        # Process cover letters if the file exists
        if os.path.isfile(cover_letters_file_path):
            try:
                self._process_cover_letters(cover_letters_file_path, output_folder)
                self._process_cvs(cvs_file_path, output_folder)
                print("Cover letters organized successfully.")
            except Exception as e:
                print(f"Error organizing cover letters: {str(e)}")

        # Process CVs if the file exists
        if os.path.isfile(cvs_file_path):
            try:
                self._process_cvs(cvs_file_path, output_folder)
                print("CVs organized successfully.")
            except Exception as e:
                print(f"Error organizing CVs: {str(e)}")

        return output # Always return the original output
    
    def _process_cover_letters(self, file_path, output_folder):
        """
        Process the cover letters JSON file and create directories for each company-job title pair.
        Each directory will contain a 'cover_letter.docx' file.
        """
        with open(file_path, "r", encoding="utf-8") as file:
            try:
                data = json.load(file)
            except json.JSONDecodeError as e:
                raise ValueError(f"Error decoding JSON from {file_path}: {str(e)}")
            
            cover_letters = data.get("tailored_cover_letters", [])
            if not cover_letters:
                raise SystemExit("No tailored cover letters found in the output file.")
            
            for letter in cover_letters:
                company = sanitize(letter.get("company_name", "Unknown Company"))
                job_title = sanitize(letter.get("job_title", "Unknown Job Title"))

                folder = f"output/{company}-{job_title}"

                os.makedirs(folder, exist_ok=True)
                cover_letter_path = os.path.join(folder, "cover_letter.docx")
                doc = docx.Document()
                content = letter.get("cover_letter_content", "No content provided.")

                for para in content.split("\n\n"):
                    doc.add_paragraph(para.strip())

                doc.save(cover_letter_path)

    def _process_cvs(self, file_path, output_folder):
        """
        Process the CVs JSON file and create directories for each company-job title pair.
        Each directory will contain a 'cv.docx' file.
        """
        with open(file_path, "r", encoding="utf-8") as file:
            try:
                data = json.load(file)
            except json.JSONDecodeError as e:
                raise ValueError(f"Error decoding JSON from {file_path}: {str(e)}")
            
            cvs = data.get("tailored_cvs", [])
            if not cvs:
                raise SystemExit("No tailored CVs found in the output file.")
            
            for cv in cvs:
                company = sanitize(cv.get("company_name", "Unknown Company"))
                job_title = sanitize(cv.get("job_title", "Unknown Job Title"))

                folder = f"output/{company}-{job_title}"
                
                os.makedirs(folder, exist_ok=True)
                cv_path = os.path.join(folder, "cv.docx")
                doc = docx.Document()

                content = cv.get("tailored_cv_content", {})
                contact_info = content.get("contact_info", {})
                summary = content.get("summary", "")
                work_experience = content.get("work_experience", [])
                education = content.get("education", [])
                certifications = content.get("certifications", [])
                skills = content.get("skills", {})
                projects = content.get("projects", {})
                interests = content.get("interests", [])

                # Contact Info block
                doc.add_heading("Contact Information", level=1)
                for key, value in contact_info.items():
                    doc.add_paragraph(f"{key.capitalize()}: {value}")
                
                # Summary block
                doc.add_heading("Professional Summary", level=1)
                doc.add_paragraph(summary.strip())

                # Work Experience block
                doc.add_heading("Work Experience", level=1)
                for experience in work_experience:
                    doc.add_heading(experience.get("job_title", "Job Title"), level=2)
                    doc.add_paragraph(f"Company: {experience.get('company_name', 'Unknown Company')}")
                    doc.add_paragraph(f"Date: {experience.get('dates', 'N/A')}")
                    doc.add_paragraph("Responsibilities:")
                    for responsibility in experience.get("responsibilities", []):
                        doc.add_paragraph(f"{responsibility.strip()}", style='ListBullet')
                
                # Education block
                doc.add_heading("Education", level=1)
                for edu in education:
                    doc.add_heading(edu.get("degree", ""), level=2)
                    doc.add_paragraph(f"School: {edu.get('university', '')}")
                    doc.add_paragraph(f"Date: {edu.get('dates', '')}")

                # Certifications block
                doc.add_heading("Certifications", level=1)
                for cert in certifications:
                    doc.add_paragraph(f"{cert.get('name', '')}: {cert.get('score', '')}")

                # Skills block
                doc.add_heading("Skills", level=1)
                for key, value in skills.items():
                    doc.add_paragraph(f"{key.replace('_', ' ').capitalize()}: {', '.join(value)}")

                # Projects block
                doc.add_heading("Projects", level=1)
                for project in projects:
                    doc.add_heading(project.get("project_name", ""), level=2)
                    doc.add_paragraph(project.get("description", ""))
                    doc.add_paragraph(f"Link: {project.get('link', '')}")

                # Interests block
                doc.add_heading("Interests", level=1)
                for interest in interests:
                    doc.add_paragraph(f"{interest.strip()}", style='ListBullet')

                doc.save(cv_path)

    # -- Define agents --
    @agent
    def cv_parser(self) -> Agent:
        return Agent(
            config=self.agents_config['cv_parser'],
            verbose=True,
            max_rpm=MAX_RPM,
            llm=gemini_llm,
            embedder={
                "provider": "google",
                "config": {
                    "model": "models/text-embedding-004",
                    "api_key": GEMINI_API_KEY,
                }
            }
        )

    @agent
    def job_scout(self) -> Agent:
        return Agent(
            config=self.agents_config['job_scout'],
            verbose=True,
            llm=gemini_llm,
            max_rpm=MAX_RPM,
            tools=[search_tool, scrape_tool]
        )
    
    @agent
    def cv_tailor(self) -> Agent:
        return Agent(
            config=self.agents_config['cv_tailor'],
            verbose=True,
            llm=gemini_llm,
            max_rpm=MAX_RPM,
        )

    @agent
    def cover_letter_writer(self) -> Agent:
        return Agent(
            config=self.agents_config['cover_letter_writer'],
            verbose=True,
            llm=gemini_llm,
            max_rpm=MAX_RPM,
        )

    # -- Define tasks --
    @task
    def parse_cv_task(self) -> Task:
        return Task(
            config=self.tasks_config['parse_cv_task'],
            output_json=Resume,
            output_file="output/structured_resume.json",
        )

    @task
    def search_jobs_task(self) -> Task:
        return Task(
            config=self.tasks_config['search_jobs_task'],
            output_json=JobPostings,
            output_file="output/job_postings.json"
        )
    
    @task
    def tailor_cv_task(self) -> Task:
        return Task(
            config=self.tasks_config['tailor_cv_task'],
            context=[self.parse_cv_task(), self.search_jobs_task()],
            output_json=TailoredCVs,
            output_file="output/tailored_cv.json",
        )

    @task
    def write_cover_letter_task(self) -> Task:
        return Task(
            config=self.tasks_config['write_cover_letter_task'],
            context=[self.parse_cv_task(), self.search_jobs_task()],
            async_execution=True,
            output_json=TailoredCoverLetters,
            output_file="output/tailored_cover_letter.json",
        )

    @crew
    def crew(self) -> Crew:
        """Creates the JobSearchAgent crew"""

        return Crew(
            agents=self.agents, # Automatically created by the @agent decorator
            tasks=self.tasks, # Automatically created by the @task decorator
            process=Process.sequential,
            verbose=True,
        )
