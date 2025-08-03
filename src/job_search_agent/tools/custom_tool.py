from crewai.tools import BaseTool
from typing import Type
from pydantic import BaseModel, Field
import os, json, re
import docx, pdfplumber


class MyCustomToolInput(BaseModel):
    """Input schema for MyCustomTool."""
    argument: str = Field(..., description="File path to the CV to be parsed.")

class cvParserTool(BaseTool):
    name: str = "CV Parser Tool"
    description: str = (
        "Parses a CV file in .pdf, .docx, .txt, or .md format and returns its text content."
    )
    args_schema: Type[BaseModel] = MyCustomToolInput

    def _run(self, argument: str) -> str:
        # Check if the file exists and extract its extension
        if not os.path.exists(argument):
            raise FileNotFoundError(f"The file {argument} does not exist.")
        ext = os.path.splitext(argument)[1].lower()

        # Parse the file based on its extension
        try:
            if ext == ".pdf":
                return self._parse_pdf(argument)
            elif ext == ".docx":
                return self._parse_docx(argument)
            elif ext in [".txt", ".md"]:
                return self._parse_text(argument)
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            raise RuntimeError(f"Error parsing file {argument}: {str(e)}")
        
    def _parse_pdf(self, file_path: str) -> str:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _parse_docx(self, file_path: str) -> str:
        text = ""
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text.strip()

    def _parse_text(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
        
def check_resume_and_extract_content(input_folder: str = "input") -> str:
    """
    Verify the input file path exists and is accessible.
    If valid, extract the content and return it in plain text format.
    If invalid, raise an error.
    """
    print(f"Checking resume in folder: {input_folder}")
    allowed_ext = ['.pdf', '.docx', '.txt', '.md']
    resume_content = ""
    resume_found = False

    for filename in os.listdir(input_folder):
        ext = os.path.splitext(filename)[1].lower()
        file_path = os.path.join(input_folder, filename)
        if os.path.isfile(file_path) and ext in allowed_ext:
            resume_found = True
            print(f"\n📃 Parsing file: {file_path}")

            try:
                if ext == ".pdf":
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            resume_content += page.extract_text() + "\n"
                elif ext == ".docx":
                    doc = docx.Document(file_path)
                    for para in doc.paragraphs:
                        resume_content += para.text + "\n"
                elif ext in [".txt", ".md"]:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        resume_content += file.read() + "\n"
                else:
                    raise ValueError(f"Unsupported file format: {ext}")
                break  # Exit loop after first valid file
            except Exception as e:
                raise RuntimeError(f"Error parsing file {file_path}: {str(e)}")
    
    if not resume_found:
        raise SystemExit("No valid resume file found in the input folder.")
    if not resume_content.strip():
        raise SystemExit("The resume file is empty or could not be read.")
    print(f"Successfully parsed file: {file_path}")
    return resume_content.strip()

def sanitize(text):
    text = re.sub(r'[\\/*?:"<>|]', "", text)  # remove invalid chars
    text = text.strip().replace(" ", "_")     # optional: replace spaces
    return text

def organize_output_files(output_folder: str = "output"):
    """
    Read output/tailored_cv.json & output/tailored_cover_letter.json files.
    Separate the content by company and job title into subdirectories named '{company}-{title}'.
    Each subdirectory will contain a 'cv.docx' and 'cover_letter.docx' file matching the company and job title.
    """
    print("\n🗂️ Organizing output files...")

    # Ensure output folder and tailored_cv.md & tailored_cover_letter.md files exist
    if not os.path.exists(output_folder):
        os.mkdir(output_folder)
    cvs_file_path = os.path.join(output_folder, "tailored_cv.json")
    cover_letters_file_path = os.path.join(output_folder, "tailored_cover_letter.json")

    if not os.path.isfile(cvs_file_path) or not os.path.isfile(cover_letters_file_path):
        raise SystemExit("Required output files do not exist.")
    
    with open(cover_letters_file_path, "r", encoding="utf-8") as file:
        data = json.load(file)
        cover_letters = data.get("tailored_cover_letters", [])

        if not cover_letters:
            raise SystemExit("No tailored cover letters found in the output file.")
        
        for letter in cover_letters:
            company = sanitize(letter.get("company_name", "Unknown Company"))
            job_title = sanitize(letter.get("job_title", "Unknown Job Title"))

            folder = f"output/{company}-{job_title}"

            os.makedirs(folder, exist_ok=True)
            cover_letter_path = os.path.join(folder, "cover_letter.docx")
            doc = docx.Document()
            content = letter.get("cover_letter_content", "No content provided.")

            for para in content.split("\n\n"):
                doc.add_paragraph(para.strip())

            doc.save(cover_letter_path)


    with open(cvs_file_path, "r", encoding="utf-8") as file:
        data = json.load(file)
        cvs = data.get("tailored_cvs", [])

        if not cvs:
            raise SystemExit("No tailored CVs found in the output file.")
        
        for cv in cvs:
            company = sanitize(cv.get("company_name", "Unknown Company"))
            job_title = sanitize(cv.get("job_title", "Unknown Job Title"))

            folder = f"output/{company}-{job_title}"
            
            os.makedirs(folder, exist_ok=True)
            cv_path = os.path.join(folder, "cv.docx")
            doc = docx.Document()

            content = cv.get("tailored_cv_content", {})
            contact_info = content.get("contact_info", {})
            summary = content.get("summary", "")
            work_experience = content.get("work_experience", [])
            education = content.get("education", [])
            certifications = content.get("certifications", [])
            skills = content.get("skills", {})
            projects = content.get("projects", {})
            interests = content.get("interests", [])

            # Contact Info block
            doc.add_heading("Contact Information", level=1)
            for key, value in contact_info.items():
                doc.add_paragraph(f"{key.capitalize()}: {value}")
            
            # Summary block
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary.strip())

            # Work Experience block
            doc.add_heading("Work Experience", level=1)
            for experience in work_experience:
                doc.add_heading(experience.get("job_title", "Job Title"), level=2)
                doc.add_paragraph(f"Company: {experience.get('company_name', 'Unknown Company')}")
                doc.add_paragraph(f"Date: {experience.get('dates', 'N/A')}")
                doc.add_paragraph("Responsibilities:")
                for responsibility in experience.get("responsibilities", []):
                    doc.add_paragraph(f"{responsibility.strip()}", style='ListBullet')
            
            # Education block
            doc.add_heading("Education", level=1)
            for edu in education:
                doc.add_heading(edu.get("degree", ""), level=2)
                doc.add_paragraph(f"School: {edu.get('university', '')}")
                doc.add_paragraph(f"Date: {edu.get('dates', '')}")

            # Certifications block
            doc.add_heading("Certifications", level=1)
            for cert in certifications:
                doc.add_paragraph(f"{cert.get('name', '')}: {cert.get('score', '')}")

            # Skills block
            doc.add_heading("Skills", level=1)
            for key, value in skills.items():
                doc.add_paragraph(f"{key.replace('_', ' ').capitalize()}: {', '.join(value)}")

            # Projects block
            doc.add_heading("Projects", level=1)
            for project in projects:
                doc.add_heading(project.get("project_name", ""), level=2)
                doc.add_paragraph(project.get("description", ""))
                doc.add_paragraph(f"Link: {project.get('link', '')}")

            # Interests block
            doc.add_heading("Interests", level=1)
            for interest in interests:
                doc.add_paragraph(f"{interest.strip()}", style='ListBullet')

            doc.save(cv_path)